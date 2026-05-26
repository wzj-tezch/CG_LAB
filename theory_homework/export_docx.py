# -*- coding: utf-8 -*-
"""将 theory_homework 下的 Markdown 答案批量转为 Word。"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt

BASE = os.path.dirname(os.path.abspath(__file__))
HW_DIRS = ["hw1_basics", "hw2_geometry", "hw3_rendering", "hw4_animation"]


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def md_to_docx(md_path, out_path):
    doc = Document()
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    in_code = False
    code_buf = []
    md_dir = os.path.dirname(md_path)

    for raw in lines:
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        if not line.strip():
            continue
        if line.strip() == "---":
            continue

        m_img = re.match(r"!\[.*?\]\((.*?)\)", line.strip())
        if m_img:
            path = os.path.join(md_dir, m_img.group(1))
            if os.path.isfile(path):
                doc.add_picture(path, width=Inches(5.5))
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("|") and re.match(r"^\|[-| :]+\|$", line.strip()):
            continue
        else:
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            text = re.sub(r"`([^`]+)`", r"\1", text)
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            text = re.sub(r"\\[\(\)]", "", text)
            if text.strip():
                doc.add_paragraph(text.strip())

    doc.save(out_path)
    print("Saved:", out_path)


def main():
    for hw in HW_DIRS:
        folder = os.path.join(BASE, hw)
        for name in os.listdir(folder):
            if name.endswith(".md"):
                md_path = os.path.join(folder, name)
                out_path = md_path[:-3] + ".docx"
                md_to_docx(md_path, out_path)


if __name__ == "__main__":
    main()
