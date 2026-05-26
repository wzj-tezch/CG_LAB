# -*- coding: utf-8 -*-
"""合并四份理论课作业为一份排版 docx，并导出 PDF。"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
OUT_DOCX = os.path.join(BASE, "202411081003-武子杰-图形学理论作业.docx")
OUT_PDF = os.path.join(BASE, "202411081003-武子杰-图形学理论作业.pdf")

STUDENT_ID = "202411081003"
STUDENT_NAME = "武子杰"

SECTIONS = [
    {
        "title": "作业一 · 基础概念部分",
        "md": os.path.join(BASE, "hw1_basics", "README.md"),
        "images": [],
    },
    {
        "title": "作业二 · 几何部分",
        "md": os.path.join(BASE, "hw2_geometry", "README.md"),
        "images": [
            (r"de_casteljau_reference\.png", os.path.join(BASE, "hw2_geometry", "de_casteljau_reference.png"), "de Casteljau 三角剖分示意图"),
        ],
    },
    {
        "title": "作业三 · 渲染部分",
        "md": os.path.join(BASE, "hw3_rendering", "README.md"),
        "images": [],
    },
    {
        "title": "作业四 · 动画部分",
        "md": os.path.join(BASE, "hw4_animation", "README.md"),
        "images": [
            (r"physics_sim_flow\.png", os.path.join(BASE, "hw4_animation", "physics_sim_flow.png"), "物理仿真计算流程图"),
            (r"skinning_lbs\.png", os.path.join(BASE, "hw4_animation", "skinning_lbs.png"), "线性混合蒙皮（LBS）示意图"),
        ],
    },
]


def set_run_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def setup_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.35
    pf.space_after = Pt(4)

    for level, size in [(1, 16), (2, 14), (3, 13)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "黑体"
        h._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        h.paragraph_format.space_before = Pt(10 if level > 1 else 14)
        h.paragraph_format.space_after = Pt(6)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("计算机图形学\n理论课作业")
    set_run_font(r, "黑体", 26, bold=True, color=RGBColor(0x1A, 0x1A, 0x2E))

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("基础知识 · 几何 · 渲染 · 动画")
    set_run_font(sr, "宋体", 14, color=RGBColor(0x55, 0x55, 0x55))

    for _ in range(3):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(f"学号：{STUDENT_ID}\n姓名：{STUDENT_NAME}")
    set_run_font(ir, "宋体", 14)

    doc.add_page_break()


def add_toc(doc):
    h = doc.add_heading("目  录", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    items = [
        "作业一  基础概念部分",
        "作业二  几何部分",
        "作业三  渲染部分",
        "作业四  动画部分",
    ]
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(f"{i}.  {item}")
        set_run_font(r, "宋体", 12)
    doc.add_page_break()


def add_rich_text(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, "宋体", 12, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 10.5)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "宋体", 12)


def add_table(doc, rows):
    if len(rows) < 2:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
            cell_text = re.sub(r"`([^`]+)`", r"\1", cell_text)
            cell = table.rows[ri].cells[ci]
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, "宋体", 10.5, bold=(ri == 0))
    doc.add_paragraph()


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    set_run_font(run, "Consolas", 9.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def try_insert_image(doc, line, image_map, md_dir):
    m = re.search(r"\[`?([^`\]]+)`?\]\(([^)]+)\)", line)
    if not m:
        return False
    ref = m.group(2)
    for pattern, path, caption in image_map:
        if re.search(pattern, ref) and os.path.isfile(path):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(5.2))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap.add_run(f"图：{caption}")
            set_run_font(cr, "宋体", 10, color=RGBColor(0x66, 0x66, 0x66))
            doc.add_paragraph()
            return True
    return False


def render_markdown(doc, md_path, section_title, image_map):
    doc.add_heading(section_title, level=1)
    md_dir = os.path.dirname(md_path)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    in_code = False
    code_buf = []
    table_buf = []
    skip_header = True

    for raw in lines:
        line = raw.rstrip("\n")

        if skip_header:
            if line.startswith("---"):
                skip_header = False
            continue

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if re.match(r"^[-| :]+$", line.strip().replace("|", "").replace("-", "").replace(":", "")):
                continue
            table_buf.append(cells)
            continue
        elif table_buf:
            add_table(doc, table_buf)
            table_buf = []

        if not line.strip():
            continue
        if line.strip() == "---":
            continue

        if try_insert_image(doc, line, image_map, md_dir):
            continue

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:].strip())
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\.\s", "", line))
        else:
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line)
            p = doc.add_paragraph()
            add_rich_text(p, text)

    if table_buf:
        add_table(doc, table_buf)


def build_docx():
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_toc(doc)

    for i, sec in enumerate(SECTIONS):
        render_markdown(doc, sec["md"], sec["title"], sec["images"])
        if i < len(SECTIONS) - 1:
            doc.add_page_break()

    doc.save(OUT_DOCX)
    print("Saved DOCX:", OUT_DOCX)
    return OUT_DOCX


def build_pdf(docx_path):
    from docx2pdf import convert
    convert(docx_path, OUT_PDF)
    print("Saved PDF:", OUT_PDF)


if __name__ == "__main__":
    docx_path = build_docx()
    build_pdf(docx_path)
